"""
readers/docx_reader.py
อ่าน .docx แปลงเป็น Markdown

Word มี heading style อยู่แล้ว (Heading 1, Heading 2, ...) แปลงเป็น # ## ได้ตรง ๆ
ทำให้ structure-aware chunking ทำงานได้ดีเหมือนกับไฟล์ Markdown จริง
"""

from pathlib import Path

from readers.base import Document, ReaderError, escape_md_cell, resolve_path


def _style_to_markdown(style_name: str, text: str) -> str:
    """แปลง Word style เป็น Markdown"""
    s = (style_name or "").lower()

    if s.startswith("heading"):
        # "Heading 2" -> ##  (จำกัดที่ 6 ตามมาตรฐาน Markdown)
        digits = "".join(c for c in s if c.isdigit())
        level = min(int(digits), 6) if digits else 1
        return f"{'#' * level} {text}"

    if "list" in s:
        return f"- {text}"
    if "quote" in s:
        return f"> {text}"
    return text


def read(path: str | Path) -> Document:
    p = resolve_path(path)

    try:
        import docx
    except ImportError:
        raise ReaderError("ต้องติดตั้ง python-docx ก่อน: pip install python-docx") from None

    try:
        document = docx.Document(str(p))
    except Exception as e:
        raise ReaderError(f"เปิดไฟล์ไม่ได้: {p.name} ({e})") from e

    parts = []
    has_heading = False

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        converted = _style_to_markdown(para.style.name if para.style else "", text)
        if converted.startswith("#"):
            has_heading = True
            parts.append("")  # เว้นบรรทัดก่อนหัวข้อ
        parts.append(converted)

    # ตารางใน Word -> ตาราง Markdown
    for idx, table in enumerate(document.tables, start=1):
        rows = [[escape_md_cell(c.text) for c in r.cells] for r in table.rows]
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            continue

        parts.append("")
        parts.append(f"## ตารางที่ {idx}")
        parts.append("")
        header, body = rows[0], rows[1:]
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join("---" for _ in header) + " |")
        for r in body:
            cells = [r[i] if i < len(r) else "" for i in range(len(header))]
            parts.append("| " + " | ".join(cells) + " |")
        parts.append("")
        has_heading = True

    text = "\n".join(parts).strip()

    # ไม่มี heading เลย -> ใส่ชื่อไฟล์เป็น H1 ให้ chunking มีจุดอ้างอิง
    if text and not has_heading:
        text = f"# {p.stem}\n\n{text}"

    # นับรูปที่ฝังอยู่ เพื่อให้รู้ว่าเอกสารนี้มีเนื้อหาที่อ่านไม่ได้มากแค่ไหน
    n_images = sum(1 for r in document.part.rels.values() if "image" in r.reltype)
    if n_images:
        parts.append("")
        parts.append(f"*(เอกสารนี้มีรูปภาพ {n_images} รูป ซึ่งยังไม่ได้อ่านเนื้อหาข้างใน)*")
        text = "\n".join(parts).strip()
        if not has_heading:
            text = f"# {p.stem}\n\n{text}"

    doc = Document(
        text=text,
        source_path=str(p),
        file_type="docx",
        metadata={
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "images": n_images,
        },
    )

    if doc.is_empty:
        raise ReaderError(f"ไม่มีเนื้อหาข้อความใน {p.name}")

    return doc
